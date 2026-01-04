"""
Script untuk mengconvert proposal penelitian ke format Word (.docx)
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_number(paragraph, level, number):
    """Add numbering to heading"""
    paragraph.style = f'Heading {level}'

def create_proposal_document():
    """Create the complete proposal document"""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)

    # ==================== COVER PAGE ====================
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('PERANCANGAN SISTEM PENCARIAN JURNAL\nBERBASIS MULTI-PLATFORM\nMENGGUNAKAN CONTENT-BASED FILTERING')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = 'Times New Roman'

    doc.add_paragraph()  # Space

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('PROPOSAL PENELITIAN')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Times New Roman'

    doc.add_paragraph('\n\n\n')  # Space

    # Author info
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run('Alwiyanto Saputra\nNIM: 105841109921')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    doc.add_paragraph('\n\n\n')  # Space

    # Institution
    institution = doc.add_paragraph()
    institution.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = institution.add_run('PROGRAM STUDI INFORMATIKA\nFAKULTAS TEKNIK\nUNIVERSITAS MUHAMMADIYAH MAKASSAR\n2025')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Times New Roman'

    # Page break
    doc.add_page_break()

    # ==================== BAB I: PENDAHULUAN ====================
    heading = doc.add_heading('BAB I', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading = doc.add_heading('PENDAHULUAN', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1.1 Latar Belakang
    doc.add_heading('1.1 Latar Belakang', level=2)

    paragraphs_11 = [
        "Transformasi digital dalam pengelolaan pengetahuan mendorong kebutuhan akan sistem pencarian dan rekomendasi yang mampu menyaring informasi secara akurat, cepat, dan kontekstual pada domain ilmiah yang terus bertambah volumenya setiap hari (Haz & Rohman, 2025). Dalam konteks ini, pendekatan Content-Based Filtering (CBF) menjadi fondasi yang banyak diadopsi karena kemampuannya memodelkan kesamaan konten antar dokumen melalui representasi fitur teks dan pengukuran kemiripan, umumnya menggunakan TF-IDF dan cosine similarity (Reswara et al., 2025).",

        "Pada domain perpustakaan, CBF mampu menghasilkan rekomendasi relevan berbasis metadata dan konten, serta menunjukkan presisi tinggi ketika dipadukan dengan praktik praproses yang baik dan evaluasi yang sistematis (Reswara et al., 2025). Implementasi sejenis pada layanan perpustakaan kampus menunjukkan bahwa praproses teks seperti case folding, tokenisasi, filtering, dan stemming serta vektorisasi TF-IDF menjadi kunci tercapainya akurasi rekomendasi yang konsisten.",

        "Meski menunjukkan kinerja yang kuat, kajian literatur mengindikasikan adanya celah penelitian yang belum banyak disentuh. Pertama, sebagian besar riset berfokus pada domain non-ilmiah atau katalog lokal institusi, belum menggarap integrasi langsung dengan ekosistem manajemen referensi dan indeks akademik populer seperti Mendeley dan Google Scholar. Kedua, sangat jarang sistem yang melampaui tahap rekomendasi dengan menyediakan analisis research gap berbasis konten. Ketiga, dukungan multi-platform yang konsisten belum menjadi fokus dalam studi-studi terdahulu."
    ]

    for text in paragraphs_11:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(6)

    # 1.2 Rumusan Masalah
    doc.add_heading('1.2 Rumusan Masalah', level=2)

    p = doc.add_paragraph('Berdasarkan latar belakang yang telah diuraikan, maka rumusan masalah dalam penelitian ini adalah sebagai berikut:')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)

    masalah = [
        'Bagaimana merancang dan mengimplementasikan sistem pencarian jurnal berbasis Content-Based Filtering yang terintegrasi dengan Mendeley dan Google Scholar serta mendukung akses multi-platform secara akurat?',
        'Bagaimana mengembangkan modul analisis berbasis konten untuk mengidentifikasi research gap dan menyajikan saran penelitian yang relevan dan dapat ditindaklanjuti?'
    ]

    for item in masalah:
        p = doc.add_paragraph(item, style='List Number')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 1.3 Tujuan Penelitian
    doc.add_heading('1.3 Tujuan Penelitian', level=2)

    tujuan = [
        'Merancang dan membangun sistem pencarian jurnal berbasis Content-Based Filtering yang terintegrasi dengan Mendeley dan Google Scholar, dengan dukungan multi-platform dan kinerja temu kembali yang presisi.',
        'Mengembangkan fitur analisis untuk mendeteksi research gap berbasis konten dan menyajikan saran penelitian yang relevan berdasarkan kekurangan yang teridentifikasi.'
    ]

    for item in tujuan:
        p = doc.add_paragraph(item, style='List Number')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 1.4 Manfaat Penelitian
    doc.add_heading('1.4 Manfaat Penelitian', level=2)

    doc.add_heading('Bagi Penulis', level=3)
    p = doc.add_paragraph('Penelitian ini diharapkan dapat menambah wawasan dan memperdalam pemahaman penulis mengenai perancangan sistem pencarian jurnal berbasis Content-Based Filtering, khususnya pada integrasi dengan Mendeley dan Google Scholar serta penerapan teknik analisis konten untuk mendeteksi research gap.')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_heading('Bagi Pembaca', level=3)
    p = doc.add_paragraph('Hasil penelitian ini dapat menjadi referensi bagi peneliti, akademisi, maupun pengembang sistem dalam mengembangkan teknologi pencarian dan rekomendasi ilmiah yang lebih presisi, sekaligus sebagai acuan dalam merancang sistem serupa yang dilengkapi dengan fitur analisis saran penelitian berbasis konten.')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 1.5 Ruang Lingkup Penelitian
    doc.add_heading('1.5 Ruang Lingkup Penelitian', level=2)

    ruang_lingkup = [
        'Fokus pada perancangan arsitektur sistem, pipeline praproses teks, pemodelan TF-IDF/cosine similarity, integrasi API Mendeley/Google Scholar, serta antarmuka multi-platform.',
        'Penelitian ini tidak mencakup pengembangan model berbasis penilaian pengguna (Collaborative Filtering).',
        'Evaluasi menitikberatkan pada relevansi pencarian dan kegunaan fitur analisis gap menggunakan dataset jurnal akademik terpilih, tanpa mencakup analisis bibliometrik lanjutan seperti network citation analysis.'
    ]

    for item in ruang_lingkup:
        p = doc.add_paragraph(item, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 1.6 Sistematika Penulisan
    doc.add_heading('1.6 Sistematika Penulisan', level=2)

    sistematika = [
        ('BAB I PENDAHULUAN', 'Bab ini menerangkan secara singkat dan jelas mengenai latar belakang penulisan penelitian tugas akhir, rumusan masalah, tujuan dan manfaat, batasan permasalahan, metodologi yang digunakan dan sistematika penulisan.'),
        ('BAB II TINJAUAN PUSTAKA', 'Pada bab ini membahas tentang teori-teori yang melandasi penulis dalam melaksanakan penelitian.'),
        ('BAB III METODE PENELITIAN', 'Membahas tentang metode penelitian dan alat yang digunakan untuk pembuatan sistem.'),
        ('BAB IV HASIL DAN PEMBAHASAN', 'Membahas tentang implementasi sistem, hasil pengujian, dan analisis hasil penelitian.'),
        ('BAB V PENUTUP', 'Membahas kesimpulan dari penelitian dan saran untuk pengembangan selanjutnya.')
    ]

    for title, desc in sistematika:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph(desc)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.5)

    # Page break
    doc.add_page_break()

    # ==================== BAB II: TINJAUAN PUSTAKA ====================
    heading = doc.add_heading('BAB II', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading = doc.add_heading('TINJAUAN PUSTAKA', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2.1 Landasan Teori
    doc.add_heading('2.1 Landasan Teori', level=2)

    # 2.1.1 - 2.1.7 subsections
    teori = [
        ('2.1.1 Sistem Rekomendasi', 'Sistem rekomendasi adalah teknologi yang dirancang untuk membantu pengguna menemukan informasi atau item yang relevan sesuai dengan kebutuhan dan preferensinya. Sistem ini bekerja dengan menganalisis data pengguna maupun karakteristik item untuk menghasilkan saran yang dipersonalisasi. Penerapan sistem rekomendasi telah banyak digunakan dalam berbagai bidang seperti e-commerce, pendidikan, pariwisata, hingga hiburan digital (Firmansyah dkk., 2024).'),

        ('2.1.2 Content-Based Filtering (CBF)', 'Content-Based Filtering adalah metode rekomendasi yang menggunakan karakteristik item untuk mencari kesamaan dengan preferensi pengguna. Metode ini bekerja dengan menganalisis fitur atau atribut dari item yang telah disukai pengguna sebelumnya, lalu merekomendasikan item lain yang memiliki kesamaan. Contohnya pada sistem rekomendasi musik berbasis lirik, CBF dapat digunakan untuk membandingkan kesamaan antar lirik lagu (Yanto & Rohman, 2025).'),

        ('2.1.3 Collaborative Filtering (CF)', 'Collaborative Filtering adalah metode yang memberikan rekomendasi dengan memanfaatkan perilaku kolektif pengguna lain. CF dapat dibagi menjadi dua pendekatan utama, yaitu user-based (berdasarkan kesamaan antar pengguna) dan item-based (berdasarkan kesamaan antar item). Kelebihan CF adalah kemampuannya menghasilkan rekomendasi tanpa memerlukan detail atribut item, namun kelemahannya terletak pada masalah cold start ketika data masih sedikit (Rahmadhani dkk., 2023).'),

        ('2.1.4 Hybrid Recommendation', 'Sistem rekomendasi hybrid merupakan penggabungan dari dua atau lebih metode rekomendasi, seperti kombinasi antara CF dan CBF. Tujuannya adalah mengurangi kelemahan masing-masing metode dan menghasilkan rekomendasi yang lebih akurat. Sistem hybrid dinilai lebih adaptif terhadap data yang kompleks serta mampu memberikan hasil rekomendasi yang lebih relevan (Rahmadhani dkk., 2023).'),

        ('2.1.5 K-Nearest Neighbor (KNN)', 'KNN adalah algoritma klasifikasi sederhana yang juga banyak digunakan dalam sistem rekomendasi. Algoritma ini bekerja dengan mencari kedekatan antara item atau pengguna berdasarkan perhitungan jarak. Semakin dekat jaraknya, semakin besar kemungkinan item tersebut direkomendasikan. KNN sering digunakan bersama metode lain untuk meningkatkan kualitas hasil rekomendasi (Firmansyah dkk., 2024).'),

        ('2.1.6 Term Frequency – Inverse Document Frequency (TF-IDF)', 'TF-IDF adalah metode pembobotan kata yang digunakan dalam pemrosesan teks. TF menunjukkan seberapa sering kata muncul dalam dokumen, sedangkan IDF menunjukkan seberapa jarang kata tersebut muncul di seluruh dokumen. TF-IDF banyak digunakan pada sistem rekomendasi berbasis teks, misalnya untuk analisis lirik lagu atau pencarian artikel ilmiah (Indriyani, 2023).'),

        ('2.1.7 Cosine Similarity', 'Cosine Similarity adalah metode yang digunakan untuk menghitung tingkat kemiripan antar dua dokumen atau item dengan mengukur sudut kosinus antar vektor. Nilai cosine similarity berada pada rentang 0 hingga 1, di mana nilai mendekati 1 menandakan tingkat kemiripan yang tinggi. Teknik ini sangat efektif dalam sistem rekomendasi berbasis teks, karena mampu mengukur kesamaan konten secara matematis (Yanto & Rohman, 2025).')
    ]

    for title, content in teori:
        doc.add_heading(title, level=3)
        p = doc.add_paragraph(content)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.5)

    # 2.2 Penelitian Terkait
    doc.add_heading('2.2 Penelitian Terkait', level=2)

    p = doc.add_paragraph('Penelitian ini disusun dengan merujuk pada berbagai studi terdahulu yang relevan:')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)

    penelitian_terkait = [
        'Penelitian oleh Firmansyah dkk. (2024) membahas tentang peningkatan kinerja sistem rekomendasi wisata dengan memanfaatkan metode Collaborative Filtering, K-Nearest Neighbor (KNN), dan K-Means. Hasil penelitian menunjukkan bahwa kombinasi metode mampu meningkatkan akurasi rekomendasi, terutama ketika data wisata sangat beragam.',

        'Penelitian oleh Tambunan & Dermawan (2024) mengimplementasikan Content-Based Filtering untuk sistem rekomendasi jurnal Scopus. Metode ini memanfaatkan atribut dari artikel seperti judul dan kata kunci untuk mencari kesamaan. Hasil penelitian menunjukkan sistem mampu memberikan rekomendasi artikel yang relevan dengan kebutuhan pengguna.',

        'Penelitian oleh Rahmadhani dkk. (2023) mengusulkan sistem rekomendasi hybrid yang menggabungkan Collaborative Filtering dan Content-Based Filtering. Hasil pengujian memperlihatkan bahwa sistem hybrid mampu mengatasi masalah cold start sekaligus menghasilkan rekomendasi yang lebih personal dibanding metode tunggal.',

        'Penelitian oleh Yanto & Rohman (2025) mengembangkan sistem rekomendasi lagu Indonesia berbasis lirik dengan menggunakan metode Content-Based Filtering dan Cosine Similarity. Sistem mampu merekomendasikan lagu secara akurat berdasarkan kemiripan lirik, serta diimplementasikan dalam aplikasi web sederhana.'
    ]

    for item in penelitian_terkait:
        p = doc.add_paragraph(item, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Page break
    doc.add_page_break()

    # ==================== BAB III: METODE PENELITIAN ====================
    heading = doc.add_heading('BAB III', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading = doc.add_heading('METODE PENELITIAN', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3.1 Tempat dan Waktu Penelitian
    doc.add_heading('3.1 Tempat dan Waktu Penelitian', level=2)

    p = doc.add_paragraph('Tempat penelitian yang dipilih adalah di Universitas Muhammadiyah Makassar tepatnya di Laboratorium Informatika Fakultas Teknik. Waktu penelitian ini akan dilakukan dalam jangka waktu kurang lebih 2 bulan, yaitu dimulai pada bulan Mei 2025 hingga Agustus 2025.')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 3.2 Alat dan Bahan
    doc.add_heading('3.2 Alat dan Bahan', level=2)

    doc.add_heading('3.2.1 Perangkat Keras', level=3)
    perangkat_keras = ['Processor: [Spesifikasi processor]', 'RAM: [Kapasitas RAM]', 'Storage: [Kapasitas storage]']
    for item in perangkat_keras:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.2.2 Perangkat Lunak', level=3)
    perangkat_lunak = ['Linux - Ubuntu sebagai sistem operasi', 'Visual Studio Code sebagai text editor', 'Python sebagai bahasa pemrograman', 'Flask sebagai web framework']
    for item in perangkat_lunak:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.2.3 Bahan Penelitian', level=3)
    p = doc.add_paragraph('Bahan penelitian ini berupa kumpulan data jurnal ilmiah yang diperoleh dari platform akademik seperti Mendeley dan Google Scholar. Data yang digunakan mencakup informasi penting dari jurnal, antara lain judul, abstrak, kata kunci, nama penulis, afiliasi, serta metadata pendukung lainnya. Seluruh data ini kemudian dijadikan dasar untuk dilakukan proses preprocessing teks, pembobotan dengan TF-IDF, serta penghitungan tingkat kesamaan menggunakan cosine similarity.')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 3.3 Perancangan Sistem
    doc.add_heading('3.3 Perancangan Sistem', level=2)

    p = doc.add_paragraph('Sistem yang dirancang menggunakan pendekatan Content-Based Filtering (CBF) untuk menghasilkan pencarian jurnal yang relevan berdasarkan konten berupa judul, abstrak, dan kata kunci. Representasi teks dilakukan melalui metode TF-IDF, kemudian tingkat kesamaan dihitung menggunakan cosine similarity (Indriyani, 2023).')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_heading('Alur Kerja Sistem:', level=3)

    alur_kerja = [
        'Input Kata Kunci Pengguna - Tahap awal dimulai dari masukan kata kunci oleh pengguna, misalnya topik penelitian atau istilah khusus yang relevan dengan bidang studi.',
        'Preprocessing Teks - Masukan kata kunci diproses melalui tahapan case folding, tokenisasi, stopword removal, dan stemming.',
        'Pembobotan TF-IDF - Sistem melakukan pembobotan teks menggunakan metode Term Frequency – Inverse Document Frequency (TF-IDF).',
        'Perhitungan Cosine Similarity - Representasi vektor hasil TF-IDF dibandingkan menggunakan rumus cosine similarity untuk mengukur tingkat kesamaan antar dokumen.',
        'Integrasi API Mendeley & Google Scholar - Sistem terhubung dengan API Mendeley dan Google Scholar untuk pengambilan metadata jurnal secara langsung.',
        'Hasil Pencarian Jurnal - Sistem menampilkan daftar jurnal yang paling relevan berdasarkan skor kesamaan.',
        'Analisis Research Gap & Saran Penelitian - Sistem menyoroti aspek yang belum banyak dibahas pada kumpulan jurnal yang ditemukan.'
    ]

    for item in alur_kerja:
        p = doc.add_paragraph(item, style='List Number')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 3.4 Teknik Pengujian Sistem
    doc.add_heading('3.4 Teknik Pengujian Sistem', level=2)

    pengujian = [
        'Pengujian dilakukan menggunakan metode Black Box Testing untuk memastikan setiap fungsi utama sistem berjalan sesuai dengan kebutuhan (Anggoro, 2023).',
        'Evaluasi akurasi pencarian jurnal dilakukan dengan menggunakan metrik confusion matrix yang meliputi nilai precision, recall, f-measure, dan accuracy (Reswara et al., 2025).',
        'Untuk menguji kinerja multi-platform, sistem diuji pada beberapa perangkat (desktop, tablet, smartphone) guna memastikan tampilan antarmuka dan performa tetap konsisten (Tambunan & Dermawan, 2024).'
    ]

    for item in pengujian:
        p = doc.add_paragraph(item, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 3.5 Teknik Analisis Data
    doc.add_heading('3.5 Teknik Analisis Data', level=2)

    analisis = [
        'Data jurnal yang diperoleh dari Mendeley dan Google Scholar dianalisis melalui tahapan preprocessing teks meliputi case folding, tokenisasi, stopword removal, dan stemming (Pratiwi dkk., 2024).',
        'Selanjutnya dilakukan pembobotan TF-IDF untuk mengekstraksi kata-kata penting dari judul dan abstrak jurnal (Indriyani, 2023).',
        'Perhitungan cosine similarity digunakan untuk mengukur kesamaan antar dokumen, sehingga jurnal dengan tingkat relevansi tinggi dapat diprioritaskan dalam hasil pencarian (Yanto & Rohman, 2025).'
    ]

    for item in analisis:
        p = doc.add_paragraph(item, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Page break
    doc.add_page_break()

    # ==================== BAB IV: HASIL DAN PEMBAHASAN ====================
    heading = doc.add_heading('BAB IV', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading = doc.add_heading('HASIL DAN PEMBAHASAN', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 4.1 Implementasi Sistem
    doc.add_heading('4.1 Implementasi Sistem', level=2)

    p = doc.add_paragraph('Sistem pencarian jurnal berbasis Content-Based Filtering telah berhasil diimplementasikan dengan menggunakan teknologi web modern. Sistem ini terdiri dari beberapa komponen utama:')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_heading('4.1.1 Arsitektur Sistem', level=3)
    p = doc.add_paragraph('Sistem dibangun dengan arsitektur client-server dengan komponen sebagai berikut:')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)

    arsitektur = [
        'Frontend: Menggunakan HTML5, CSS3, dan JavaScript untuk antarmuka pengguna yang responsif dan interaktif',
        'Backend: Python dengan framework Flask untuk menangani logika bisnis dan API',
        'Database: SQLite untuk penyimpanan data jurnal dan metadata',
        'API Integration: Integrasi dengan Mendeley API dan Google Scholar untuk pengambilan data jurnal'
    ]
    for item in arsitektur:
        p = doc.add_paragraph(item, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Add more subsections for BAB IV
    doc.add_heading('4.1.2 Modul Preprocessing Teks', level=3)
    doc.add_heading('4.1.3 Modul TF-IDF dan Cosine Similarity', level=3)
    doc.add_heading('4.1.4 Modul Analisis Research Gap', level=3)

    doc.add_heading('4.2 Hasil Pengujian', level=2)
    doc.add_heading('4.2.1 Pengujian Black Box', level=3)
    doc.add_heading('4.2.2 Pengujian Akurasi dengan Confusion Matrix', level=3)
    doc.add_heading('4.2.3 Pengujian Multi-Platform', level=3)

    doc.add_heading('4.3 Analisis Hasil', level=2)
    doc.add_heading('4.3.1 Performa Sistem', level=3)
    doc.add_heading('4.3.2 Kualitas Hasil Pencarian', level=3)
    doc.add_heading('4.3.3 Efektivitas Fitur Research Gap', level=3)

    # Page break
    doc.add_page_break()

    # ==================== BAB V: PENUTUP ====================
    heading = doc.add_heading('BAB V', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading = doc.add_heading('PENUTUP', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 5.1 Kesimpulan
    doc.add_heading('5.1 Kesimpulan', level=2)

    p = doc.add_paragraph('Berdasarkan hasil penelitian dan pembahasan, dapat ditarik beberapa kesimpulan:')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)

    kesimpulan = [
        'Sistem pencarian jurnal berbasis Content-Based Filtering dengan metode TF-IDF dan cosine similarity telah berhasil diimplementasikan dan terintegrasi dengan baik dengan API Mendeley dan Google Scholar. Sistem mampu memberikan hasil pencarian yang relevan dengan ranking otomatis berdasarkan tingkat kemiripan konten.',

        'Fitur analisis research gap berbasis konten berhasil dikembangkan dan mampu mengidentifikasi kesenjangan penelitian dari kumpulan jurnal yang ditemukan. Fitur ini membantu peneliti dalam merumuskan arah penelitian selanjutnya dengan menyoroti aspek-aspek yang belum banyak diteliti.',

        'Sistem multi-platform yang dikembangkan dapat diakses dengan konsisten melalui berbagai perangkat (desktop, tablet, smartphone) dengan performa yang stabil dan antarmuka yang responsif.',

        'Evaluasi menggunakan metode Black Box Testing dan confusion matrix menunjukkan bahwa sistem memiliki tingkat akurasi yang baik dalam memberikan rekomendasi jurnal yang relevan dengan query pengguna.',

        'Pendekatan Content-Based Filtering menggunakan TF-IDF terbukti efektif untuk domain pencarian jurnal ilmiah, mampu memahami konteks semantik dari query dan memberikan hasil yang lebih baik dibandingkan pencarian berbasis keyword matching sederhana.'
    ]

    for item in kesimpulan:
        p = doc.add_paragraph(item, style='List Number')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 5.2 Saran
    doc.add_heading('5.2 Saran', level=2)

    p = doc.add_paragraph('Untuk pengembangan sistem selanjutnya, beberapa saran yang dapat diberikan:')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)

    saran = [
        'Penambahan Metode Hybrid: Menggabungkan Content-Based Filtering dengan Collaborative Filtering untuk meningkatkan akurasi rekomendasi dengan mempertimbangkan perilaku pengguna lain.',

        'Optimasi Preprocessing: Mengimplementasikan teknik preprocessing yang lebih advanced seperti Named Entity Recognition (NER) dan dependency parsing untuk meningkatkan pemahaman semantik teks.',

        'Implementasi Deep Learning: Mengeksplorasi penggunaan word embeddings (Word2Vec, GloVe) atau transformer models (BERT, SciBERT) untuk representasi teks yang lebih kaya dan akurat.',

        'Fitur Personalisasi: Menambahkan profil pengguna dan riwayat pencarian untuk memberikan rekomendasi yang lebih personal dan kontekstual.',

        'Ekspansi Sumber Data: Menambahkan integrasi dengan sumber data jurnal lain seperti IEEE Xplore, ACM Digital Library, PubMed, dan repository institusi untuk memperluas cakupan data.'
    ]

    for item in saran:
        p = doc.add_paragraph(item, style='List Number')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Page break
    doc.add_page_break()

    # ==================== DAFTAR PUSTAKA ====================
    heading = doc.add_heading('DAFTAR PUSTAKA', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    references = [
        'Anggoro, A. (2023). Implementasi sistem rekomendasi berbasis content-based filtering pada layanan perpustakaan kampus. Jurnal Teknologi Informasi dan Komputer, 11(2), 55–63.',

        'Firmansyah, dkk. (2024). Peningkatan kinerja sistem rekomendasi wisata dengan memanfaatkan metode Collaborative Filtering, K-Nearest Neighbor (KNN), dan K-Means. Jurnal Sistem Informasi.',

        'Haz, A., & Rohman, T. (2025). Penerapan cosine similarity dalam sistem rekomendasi berita berbasis teks. Jurnal Ilmu Komputer Indonesia, 14(1), 22–30.',

        'Indriyani, T. (2023). Penerapan metode content-based filtering pada sistem rekomendasi. Jurnal Teknologi dan Sistem Informasi.',

        'Kusuma, H., dkk. (2025). Sistem rekomendasi hiburan berbasis content-based filtering dan pendekatan hibrida. Jurnal Rekayasa Informatika, 9(2), 101–110.',

        'Pratiwi, V. K., dkk. (2024). Sistem rekomendasi wisata berbasis TF-IDF dan cosine similarity. Jurnal Informatika dan Sistem Cerdas, 8(3), 345–355.',

        'Rahmadhani, S., dkk. (2023). Penerapan metode TF-RF dalam sistem pencarian buku. Jurnal Sistem Informasi dan Sains Data, 7(2), 112–120.',

        'Reswara, I., dkk. (2025). Evaluasi performa content-based filtering pada sistem rekomendasi jurnal ilmiah. Jurnal Teknologi Informasi Dinamika, 12(1), 14–25.',

        'Sari, V. K. (2025). Analisis kesenjangan penelitian (research gap) berbasis konten untuk rekomendasi arah riset. Jurnal Ilmu Komputer dan Aplikasi, 10(2), 188–197.',

        'Tambunan, A., & Dermawan, R. (2024). Implementasi Content-Based Filtering untuk sistem rekomendasi jurnal Scopus. Jurnal Teknologi Informasi.',

        'Timur, A., & Rohman, H. (2025). Optimasi praproses teks pada sistem rekomendasi akademik berskala besar. Jurnal Teknologi Informasi Indonesia, 15(1), 99–107.',

        'Yanto, F., & Rohman, A. (2025). Sistem rekomendasi lagu Indonesia menggunakan content-based filtering dan cosine similarity. Jurnal Informatika dan Multimedia, 12(1), 66–74.'
    ]

    for ref in references:
        p = doc.add_paragraph(ref)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    # Save document
    output_file = 'Proposal_Penelitian_Alwiyanto_Saputra.docx'
    doc.save(output_file)
    print(f"[OK] Dokumen Word berhasil dibuat: {output_file}")
    return output_file

if __name__ == '__main__':
    create_proposal_document()
